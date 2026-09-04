import asyncio
import json
import sqlite3
from pathlib import Path
from typing import Any, TypedDict

from docx import Document
from docx.shared import Pt
from langgraph.checkpoint.sqlite import SqliteSaver
from langgraph.graph import END, START, StateGraph
from langgraph.types import Command, interrupt
from pydantic import ValidationError

from app.api.routers.ai import MeetingSummaryResponse, redact_sensitive_text
from app.core.config import settings
from app.core.database import SessionLocal
from app.llm.deepseek import DeepSeekClient
from app.models.entities import ChatMessage, MeetingJob, MeetingRecord
from app.services.transcription import transcribe_media


class MeetingAgentState(TypedDict, total=False):
    job_id: int
    transcript: str
    filtered_transcript: str
    minutes: dict[str, Any]
    redacted_sensitive_data: bool


def update_job(job_id: int, **values: Any) -> MeetingJob:
    with SessionLocal() as db:
        job = db.get(MeetingJob, job_id)
        if not job:
            raise RuntimeError("会议任务不存在")
        for key, value in values.items():
            setattr(job, key, value)
        db.commit(); db.refresh(job)
        return job


def transcribe_node(state: MeetingAgentState) -> dict:
    with SessionLocal() as db:
        job = db.get(MeetingJob, state["job_id"])
        if not job:
            raise RuntimeError("会议任务不存在")
        if job.transcript:
            return {"transcript": job.transcript}
        job.status = "transcribing"; job.error_message = None; db.commit()
        source = Path(job.source_path)
    transcript = transcribe_media(source)
    update_job(state["job_id"], transcript=transcript, status="awaiting_transcript_review")
    return {"transcript": transcript}


def review_transcript_node(state: MeetingAgentState) -> dict:
    update_job(state["job_id"], status="awaiting_transcript_review")
    response = interrupt({"stage": "transcript_review", "job_id": state["job_id"]})
    transcript = str((response or {}).get("transcript", "")).strip()
    if len(transcript) < 20:
        raise ValueError("转写稿至少需要20个字")
    update_job(state["job_id"], transcript=transcript, status="filtering")
    return {"transcript": transcript}


def filter_node(state: MeetingAgentState) -> dict:
    redacted, was_redacted = redact_sensitive_text(state["transcript"])
    client = DeepSeekClient()
    filtered = asyncio.run(client.complete([
        {"role": "system", "content": "你是会议文字稿清理助手。删除口头语、重复句和明显无关闲聊，保留所有事实、决定、人员职责和时间信息。不得补写原文没有的内容，只返回清理后文字。"},
        {"role": "user", "content": redacted},
    ]))
    update_job(state["job_id"], filtered_transcript=filtered, status="generating_minutes")
    return {"filtered_transcript": filtered, "redacted_sensitive_data": was_redacted}


def generate_minutes_node(state: MeetingAgentState) -> dict:
    with SessionLocal() as db:
        job = db.get(MeetingJob, state["job_id"])
        meeting_type = job.meeting_type if job else "其他"
        instruction = job.instruction if job else "请整理为标准会议纪要"
    client = DeepSeekClient()
    raw = asyncio.run(client.complete([
        {"role": "system", "content": "只根据文字稿生成会议纪要JSON，不得补写事实。字段必须为title、summary、key_points、decisions、action_items；action_items每项包含task、owner、deadline，未明确时填‘未指定’。"},
        {"role": "user", "content": f"会议类型：{meeting_type}\n整理要求：{instruction}\n\n文字稿：\n{state['filtered_transcript']}"},
    ], json_output=True))
    try:
        data = json.loads(raw)
        data.update(meeting_type=meeting_type, requires_manual_review=True, redacted_sensitive_data=state.get("redacted_sensitive_data", False))
        minutes = MeetingSummaryResponse.model_validate(data).model_dump()
    except (json.JSONDecodeError, ValidationError, TypeError) as exc:
        raise RuntimeError("AI返回的会议纪要格式不完整") from exc
    update_job(state["job_id"], minutes_json=json.dumps(minutes, ensure_ascii=False), status="awaiting_minutes_review")
    return {"minutes": minutes}


def review_minutes_node(state: MeetingAgentState) -> dict:
    update_job(state["job_id"], status="awaiting_minutes_review")
    response = interrupt({"stage": "minutes_review", "job_id": state["job_id"]})
    minutes = MeetingSummaryResponse.model_validate(response).model_dump()
    return {"minutes": minutes}


def build_minutes_docx(target: Path, minutes: dict, transcript: str) -> Path:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"; normal.font.size = Pt(11)
    document.add_heading(minutes["title"], level=0)
    document.add_paragraph(f"会议类型：{minutes['meeting_type']}")
    document.add_heading("会议摘要", level=1); document.add_paragraph(minutes["summary"])
    document.add_heading("主要内容", level=1)
    for item in minutes["key_points"]:
        document.add_paragraph(item, style="List Bullet")
    document.add_heading("会议决定", level=1)
    for item in minutes["decisions"]:
        document.add_paragraph(item, style="List Bullet")
    document.add_heading("待办事项", level=1)
    table = document.add_table(rows=1, cols=3); table.style = "Table Grid"
    for cell, title in zip(table.rows[0].cells, ("任务", "负责人", "截止时间")):
        cell.text = title
    for item in minutes["action_items"]:
        cells = table.add_row().cells
        cells[0].text = item["task"]; cells[1].text = item["owner"]; cells[2].text = item["deadline"]
    document.add_heading("经确认的文字稿", level=1); document.add_paragraph(transcript)
    document.save(target)
    return target


def build_docx(job_id: int, minutes: dict, transcript: str) -> Path:
    return build_minutes_docx(settings.meeting_documents_dir / f"meeting-{job_id}.docx", minutes, transcript)


def finalize_node(state: MeetingAgentState) -> dict:
    minutes = MeetingSummaryResponse.model_validate(state["minutes"]).model_dump()
    with SessionLocal() as db:
        job = db.get(MeetingJob, state["job_id"])
        if not job:
            raise RuntimeError("会议任务不存在")
        record = db.get(MeetingRecord, job.meeting_record_id) if job.meeting_record_id else None
        if not record:
            record = MeetingRecord(
                class_id=job.class_id, author_id=job.user_id, meeting_type=job.meeting_type,
                title=minutes["title"], transcript=job.transcript, summary=minutes["summary"],
                summary_json=json.dumps({"summary": minutes["summary"]}, ensure_ascii=False),
                key_points_json=json.dumps(minutes["key_points"], ensure_ascii=False),
                decisions_json=json.dumps(minutes["decisions"], ensure_ascii=False),
                action_items_json=json.dumps(minutes["action_items"], ensure_ascii=False),
                source_path=job.source_path, source_name=job.source_name,
            )
            db.add(record); db.flush(); job.meeting_record_id = record.id
        document_path = build_docx(job.id, minutes, job.transcript)
        job.minutes_json = json.dumps(minutes, ensure_ascii=False); job.document_path = str(document_path); job.status = "complete"; job.error_message = None
        message = db.query(ChatMessage).filter(ChatMessage.meeting_job_id == job.id, ChatMessage.role == "assistant").first()
        if message:
            key_points = "\n".join(f"- {item}" for item in minutes["key_points"]) or "- 无"
            decisions = "\n".join(f"- {item}" for item in minutes["decisions"]) or "- 无"
            actions = "\n".join(
                f"- {item['task']}（负责人：{item['owner']}；截止时间：{item['deadline']}）"
                for item in minutes["action_items"]
            ) or "- 无"
            message.content = (
                f"# {minutes['title']}\n\n"
                f"## 会议摘要\n{minutes['summary']}\n\n"
                f"## 主要内容\n{key_points}\n\n"
                f"## 会议决定\n{decisions}\n\n"
                f"## 待办事项\n{actions}\n\n"
                "会议纪要已经人工确认，可以下载 Word 文档。"
            )
            message.status = "complete"
        db.commit()
    return {}


def build_graph():
    builder = StateGraph(MeetingAgentState)
    builder.add_node("transcribe", transcribe_node)
    builder.add_node("review_transcript", review_transcript_node)
    builder.add_node("filter", filter_node)
    builder.add_node("generate_minutes", generate_minutes_node)
    builder.add_node("review_minutes", review_minutes_node)
    builder.add_node("finalize", finalize_node)
    builder.add_edge(START, "transcribe")
    builder.add_edge("transcribe", "review_transcript")
    builder.add_edge("review_transcript", "filter")
    builder.add_edge("filter", "generate_minutes")
    builder.add_edge("generate_minutes", "review_minutes")
    builder.add_edge("review_minutes", "finalize")
    builder.add_edge("finalize", END)
    connection = sqlite3.connect(settings.langgraph_checkpoint_path, check_same_thread=False)
    saver = SqliteSaver(connection); saver.setup()
    return builder.compile(checkpointer=saver)


meeting_graph = build_graph()


def run_meeting_graph(job_id: int, resume: dict | None = None, retry: bool = False) -> None:
    config = {"configurable": {"thread_id": f"meeting-{job_id}"}}
    try:
        if resume is not None:
            meeting_graph.invoke(Command(resume=resume), config=config)
        elif retry:
            meeting_graph.invoke(None, config=config)
        else:
            meeting_graph.invoke({"job_id": job_id}, config=config)
    except Exception as exc:
        update_job(job_id, status="failed", error_message=str(exc)[:1000])
